"""
Updates docs/Abstract.docx and docs/documentation.docx
  1. Rewrites Abstract completely (short doc, full rewrite)
  2. Fixes documentation.docx via targeted text replacements:
       - stock -> crypto throughout
       - MAPE -> DA (Direction Accuracy)
       - ten indicators -> eleven
       - wrong project name ATI -> AlgoTrade AI
       - stock ticker AAPL -> crypto ticker BTC-USD
       - keywords updated
       - RMSE + MAPE metrics -> RMSE + DA metrics

Run: python update_docs.py
Output: docs/Abstract_updated.docx, docs/documentation_updated.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# ── Global text replacement map ──────────────────────────────────────────────
REPLACEMENTS = [
    # Project/domain
    ("next-day stock returns",          "next-day cryptocurrency returns"),
    ("stock market trends",             "cryptocurrency market trends"),
    ("stock market behavior",           "cryptocurrency market behavior"),
    ("stock market prediction",         "cryptocurrency market prediction"),
    ("Stock Market Prediction",         "Cryptocurrency Market Prediction"),
    ("stock market",                    "cryptocurrency market"),
    ("stock returns",                   "cryptocurrency returns"),
    ("stock prices",                    "cryptocurrency prices"),
    ("stock price",                     "cryptocurrency price"),
    ("stock data",                      "cryptocurrency data"),
    ("historical stock data",           "historical cryptocurrency data"),
    ("stock ticker (like AAPL for Apple)","crypto ticker (like BTC-USD for Bitcoin)"),
    ("stock ticker",                    "crypto ticker"),
    ("global stock markets",            "global cryptocurrency markets"),
    ("in stocks",                       "in cryptocurrencies"),
    ("for stocks",                      "for cryptocurrencies"),
    ("buy and sell stocks",             "buy and sell cryptocurrencies"),
    ("buying and keeping the stock",    "buying and holding the cryptocurrency"),
    # Metrics
    ("MAE, RMSE, and MAPE",             "MAE, RMSE, and DA (Direction Accuracy)"),
    ("MAE, RMSE, MAPE",                 "MAE, RMSE, DA (Direction Accuracy)"),
    ("RMSE, MAPE",                      "RMSE, DA (Direction Accuracy)"),
    ("MAPE, which quantifies errors relative to actual returns",
     "DA (Direction Accuracy), which measures how often the model correctly predicted price direction"),
    ("Mean Absolute Percentage Error (MAPE, error as a percentage)",
     "Direction Accuracy (DA, % of correct up/down predictions)"),
    ("Mean Absolute Percentage Error (MAPE)",
     "Direction Accuracy — DA (% of days direction predicted correctly)"),
    ("MAPE",                            "DA (Direction Accuracy)"),
    # Indicators count
    ("ten key technical indicators",    "eleven key technical indicators"),
    ("a wide range of",                 "eleven"),
    # Wrong project name
    ("Automated Trading Intelligence (ATI)", "AlgoTrade AI — MarketPulse Optimizer"),
    ("ATI incorporates",                "AlgoTrade AI incorporates"),
    ("ATI is",                          "AlgoTrade AI is"),
    ("ATI develops",                    "AlgoTrade AI develops"),
    ("Core to ATI",                     "Core to AlgoTrade AI"),
    ("ATI involves",                    "AlgoTrade AI involves"),
    ("benefits of ATI",                 "benefits of AlgoTrade AI"),
    # Buy/Sell → Buy/Sell/Hold
    ("buy if predicted",                "buy/sell/hold if predicted"),
    ('"buy if MACD',                    '"buy/sell/hold if MACD'),
    # Signals
    ("Buy and Sell",                    "Buy, Sell, and Hold"),
    ("buy-and-hold",                    "buy-and-hold (passive baseline)"),
    # Keywords
    ("Stock Market Prediction, Backtesting",
     "Cryptocurrency Market Prediction, Bitcoin, Ethereum, Backtesting, "
     "Direction Accuracy, Sharpe Ratio"),
]


def apply_replacements(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def fix_paragraph(para):
    """Apply replacements across all runs in a paragraph, preserving formatting."""
    # Collect full text and check if replacement needed
    full = "".join(r.text for r in para.runs)
    fixed = apply_replacements(full)
    if fixed == full:
        return
    # Put all fixed text in first run, clear the rest
    if para.runs:
        para.runs[0].text = fixed
        for r in para.runs[1:]:
            r.text = ""


def fix_document(doc):
    for para in doc.paragraphs:
        fix_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_paragraph(para)


# ════════════════════════════════════════════════════════════════════════════
# 1. ABSTRACT — full rewrite
# ════════════════════════════════════════════════════════════════════════════

doc_abs = Document('docs/Abstract.docx')

# Clear all paragraphs content
for para in doc_abs.paragraphs:
    for run in para.runs:
        run.text = ""

# Rebuild paragraph by paragraph
paras = doc_abs.paragraphs

def set_para(para, text, bold=False, size=None, align=None):
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align:
        para.alignment = align

# Paragraph 0 — heading "Abstract"
if len(paras) > 0:
    set_para(paras[0], "Abstract", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER)

# Paragraph 1 — main abstract body (paragraph 1)
NEW_ABSTRACT_BODY = (
    "The rapid expansion of cryptocurrency markets has created an urgent need for "
    "intelligent systems capable of processing high-frequency market data and generating "
    "reliable, data-driven trading signals. AlgoTrade AI — MarketPulse Optimizer is an "
    "end-to-end algorithmic trading framework that integrates machine learning and deep "
    "learning to predict next-day cryptocurrency price direction and simulate profitable "
    "trading strategies. The system retrieves live price data from Yahoo Finance for "
    "major cryptocurrencies including Bitcoin (BTC), Ethereum (ETH), and Solana (SOL), "
    "then computes eleven key technical indicators — SMA, WMA, Momentum, Stochastic %K/%D, "
    "RSI, MACD, Williams %R, Accumulation/Distribution, and CCI — to capture market "
    "trend, momentum, volume, and oscillation signals."
)

if len(paras) > 1:
    set_para(paras[1], NEW_ABSTRACT_BODY)

# Paragraph 2 — second body paragraph
NEW_ABSTRACT_BODY2 = (
    "Three machine learning models — XGBoost, Random Forest, and LSTM — are trained to "
    "forecast next-day returns, evaluated using RMSE and Direction Accuracy (DA), where "
    "DA measures the percentage of days the model correctly predicted market direction "
    "(up or down). A value above 55% indicates genuine predictive power beyond the "
    "random 50% baseline. The system generates Buy, Sell, and Hold signals based on a "
    "0.1% confidence threshold and backtests the resulting strategy on the most recent "
    "20% of historical data, incorporating realistic transaction costs of 0.1% per trade "
    "and both long and short positions. Strategy performance is evaluated against a "
    "buy-and-hold passive baseline using Sharpe Ratio, Maximum Drawdown, and Win Rate. "
    "An interactive Streamlit dashboard presents live trading signals, model comparison "
    "via radar charts, feature importance analysis, cumulative return visualisation, and "
    "an investment calculator. AlgoTrade AI demonstrates how predictive modelling can "
    "transform cryptocurrency trading from intuition-driven guesswork into a rigorous, "
    "data-driven decision process."
)

if len(paras) > 2:
    set_para(paras[2], NEW_ABSTRACT_BODY2)

# Paragraph 3 — Keywords heading
if len(paras) > 3:
    set_para(paras[3], "Keywords", bold=True)

# Paragraph 4 — Keywords list
NEW_KEYWORDS = (
    "Algorithmic Trading, Machine Learning, Deep Learning, Financial Forecasting, "
    "XGBoost, Random Forest, LSTM, Technical Indicators, Cryptocurrency Market Prediction, "
    "Bitcoin, Ethereum, Backtesting, Direction Accuracy, Sharpe Ratio, Streamlit Dashboard, "
    "FinTech AI, Long Short Positions, Transaction Costs."
)
if len(paras) > 4:
    set_para(paras[4], NEW_KEYWORDS)

doc_abs.save('docs/Abstract_updated.docx')
print("Abstract_updated.docx saved.")


# ════════════════════════════════════════════════════════════════════════════
# 2. DOCUMENTATION — targeted fixes
# ════════════════════════════════════════════════════════════════════════════

doc_main = Document('docs/documentation.docx')
fix_document(doc_main)

# Also fix the abstract section inside documentation (paragraphs 3 & 4)
# Find paragraphs that match the old abstract text and replace
for i, para in enumerate(doc_main.paragraphs):
    full = "".join(r.text for r in para.runs)

    # Fix old abstract body paragraph 1
    if "forecast next-day stock returns" in full or "next-day stock returns" in full:
        new_text = apply_replacements(full)
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""

    # Fix keywords paragraph
    if "Stock Market Prediction, Backtesting, Streamlit Dashboard" in full:
        new_kw = (
            "Algorithmic Trading, Machine Learning, Deep Learning, Financial Forecasting, "
            "XGBoost, Random Forest, LSTM, Technical Indicators, Cryptocurrency Market Prediction, "
            "Bitcoin, Ethereum, Backtesting, Direction Accuracy, Sharpe Ratio, Streamlit Dashboard, "
            "FinTech AI, Long Short Positions, Transaction Costs."
        )
        if para.runs:
            para.runs[0].text = new_kw
            for r in para.runs[1:]:
                r.text = ""

    # Fix section 1.1 wrong project name ATI
    if "Automated Trading Intelligence" in full:
        new_text = full.replace(
            "Automated Trading Intelligence (ATI) is an innovative AI-driven project aimed at "
            "revolutionizing how financial markets are navigated. At its heart, ATI develops a "
            "fully autonomous system",
            "AlgoTrade AI — MarketPulse Optimizer is an innovative AI-driven system aimed at "
            "revolutionizing how cryptocurrency markets are navigated. At its heart, it develops "
            "a comprehensive automated system"
        )
        new_text = apply_replacements(new_text)
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""

    # Fix MAPE in metrics explanation paragraph
    if "Root Mean Squared Error (RMSE, average prediction error), Mean Absolute Percentage Error" in full:
        new_text = full.replace(
            "Root Mean Squared Error (RMSE, average prediction error), "
            "Mean Absolute Percentage Error (MAPE, error as a percentage), "
            "and Mean Absolute Error",
            "Root Mean Squared Error (RMSE, average prediction error), "
            "Direction Accuracy (DA, % of days the model correctly predicted up or down — "
            "50% is random, above 55% is genuinely predictive), "
            "and Mean Absolute Error"
        )
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""

doc_main.save('docs/documentation_updated.docx')
print("documentation_updated.docx saved.")

print("\nAll done. Summary of changes:")
print("  Abstract   : Full rewrite — crypto focus, 11 indicators, DA metric, realistic backtest")
print("  Documentation: Fixed stock->crypto, MAPE->DA, ATI->AlgoTrade AI, keywords, metrics")
